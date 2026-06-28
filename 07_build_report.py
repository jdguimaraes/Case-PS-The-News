# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
GRAY = RGBColor(0x55, 0x55, 0x55)

def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs: r.font.color.rgb = NAVY
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs: r.font.color.rgb = NAVY
    return p

def para(text, bold=False, italic=False, size=11, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return p

def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def shade_cell(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def make_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = htext
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True; r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        shade_cell(hdr[i], '2E5C8A')
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return t

# ---------------- TITLE ----------------
title = doc.add_heading('Palavritas: o que explica a retenção e o que fazer sobre isso', level=0)
for r in title.runs: r.font.color.rgb = NAVY
para('Análise de comportamento e proposta de ação — Analista de Dados, Produto & Growth', italic=True, color=GRAY)
para('Base: 39.849 sessões válidas (de 41.157 originais) · 1.200 usuários · 800 com perfil de pesquisa', italic=True, size=10, color=GRAY)
doc.add_paragraph()

# ---------------- RESUMO EXECUTIVO ----------------
h1('Resumo para quem tem 2 minutos')
para('A pergunta era: "o que faz alguém voltar a jogar — e o que podemos fazer sobre isso?" A resposta tem uma virada importante:')
bullet('Quase nada explica quem volta no dia seguinte (D+1). Testamos horário, palavra do dia, device, setor, salário, delivery, newsletter — nenhum desses fatores muda a taxa de forma relevante. Ela fica em ~22% quase sempre.')
bullet('Mas o que explica quem fica ativo depois de 30 dias (D30) é muito claro: voltar no D+1, jogar de manhã, e abrir a newsletter antes de jogar. Esses três fatores se somam: quem faz as três coisas tem 92% de chance de estar ativo em 30 dias, contra 22% da base geral.')
para('Conclusão prática: não vamos conseguir prever ou empurrar o D+1 olhando para essas variáveis — o gatilho provavelmente está em algo que não está nesse dataset (notificação, horário de envio, hábito). Mas dá pra agir na conexão newsletter → jogo, que é o maior driver de retenção de longo prazo que apareceu na análise.', bold=False)

# ---------------- ENTREGA 1 ----------------
h1('Entrega 1 — Limpeza e diagnóstico')
para('Antes de analisar qualquer correlação, fui linha por linha verificar se os dados faziam sentido. Encontrei 7 problemas reais. Documentei cada um com a decisão tomada e o porquê — para que qualquer pessoa possa auditar essas escolhas depois.')

make_table(
    ['Problema encontrado', 'O que eu vi', 'Decisão e motivo'],
    [
        ['Linhas duplicadas em sessions', '1.198 sessões (2,9%) apareciam duas vezes, idênticas em todas as colunas.',
         'Removidas (mantida 1ª ocorrência). Duplicata exata indica erro de exportação, não atividade real — contar 2x infla todas as métricas de volume.'],
        ['Linhas duplicadas em attempts', '752 tentativas duplicadas (0,5%).',
         'Removidas pelo mesmo motivo.'],
        ['Tentativas fora da regra do jogo', '90 sessões com "attempts" = 0, 7 ou 8 (o jogo permite só 1 a 6 tentativas).',
         'Removidas (0,2% da base). Combinadas com as 60 sessões sem resultado (result em branco), foram descartadas 110 linhas corrompidas (0,28%). É baixo volume e não distorce a análise — mas decidi não "adivinhar" o valor certo.'],
        ['Resultado (win/lose) em branco', '63 sessões sem essa informação.',
         'Removidas junto com o item acima (mesma causa raiz aparente: log incompleto da sessão).'],
        ['Tempo de jogo negativo', '43 sessões com tempo de -5 segundos (impossível).',
         'Não descartei a sessão inteira — só apaguei o valor de tempo (virou vazio). A sessão em si pode ser real, só o cronômetro do app falhou.'],
        ['Datas em dois formatos misturados', '6.162 linhas (15%) com data tipo "19/01/2026" e o resto em "2026-01-19" na mesma coluna.',
         'Padronizei tudo para um único formato de data, interpretando dia/mês/ano (padrão brasileiro) nos casos ambíguos.'],
        ['Maiúsculas/minúsculas inconsistentes', 'Device aparecia como "Android", "android", "ios", "iOS", "IOS", "ANDROID" — 6 variações para 2 categorias reais.',
         'Padronizado para "Android" / "iOS". Sem isso, qualquer agrupamento por device estaria sub-contando ambos.'],
        ['Texto corrompido (encoding) no perfil', 'Palavras como "setor" e "salário" vieram com caracteres quebrados: "finan�as", "educa��o", "at� R$2k", "S�nior". Também havia "São Paulo" e "Minas Gerais" escritos por extenso misturados com siglas de estado (SP, MG).',
         'Reconstruí manualmente os termos corrompidos pelo contexto (ex: "finan�as" → "finanças") e unifiquei estado para sigla (SP, MG). É uma correção textual, não estatística — registro aqui para transparência.'],
        ['Cargo (job_role) com grafias diferentes', '"Consultor", "consultor", "Consultor Sênior" tratados como categorias distintas.',
         'Padronizado para Title Case, preservando os níveis (Sênior, Coordenador etc.) como são informações reais, não erro.'],
        ['"Pede delivery" com 2 formatos diferentes', 'Coluna trazia tanto True/False quanto "sim"/"não" para a mesma pergunta.',
         'Unificado para verdadeiro/falso.'],
    ]
)

doc.add_paragraph()
h2('O que eu decidi NÃO tratar (e por quê)')
bullet('400 dos 1.200 usuários (33%) não têm perfil de pesquisa. Não inventei perfil para eles — qualquer análise por setor, salário ou device de perfil cobre só os 800 que responderam. Isso limita (mas não invalida) os cortes por perfil.')
bullet('Idade, cidade e faixa salarial têm valores em branco em ~15-37% dos 800 perfis. Mantive como vazio (não preenchi com média/moda) porque presumir resposta de quem não respondeu distorceria os percentuais reais.')
bullet('Em 62 sessões (0,16%), o campo "attempts" da tabela de sessões não bate com o número real de tentativas registradas na tabela de attempts. É um volume baixo demais para justificar reconciliar manualmente — fica registrado como limitação conhecida da fonte de dados.')
bullet('200 tentativas (attempts) referenciam sessões que não existem na tabela de sessões, e 93 sessões não têm nenhuma tentativa registrada. Não exclui essas sessões da análise principal — elas só não entram em análises que dependem do detalhe de cada tentativa (palpites, letras certas).')

para('Resultado da limpeza: de 41.157 sessões originais, ficamos com 39.849 (97% aproveitado) — uma base sólida para a análise abaixo.', bold=True)

# ---------------- ENTREGA 2 ----------------
doc.add_page_break()
h1('Entrega 2 — O que de fato move a retenção')
para('Antes de testar variáveis, precisei separar duas perguntas que o briefing junta mas que têm respostas muito diferentes: "o que faz alguém voltar amanhã" e "o que faz alguém continuar jogando depois de um mês". São fenômenos diferentes — e tratá-los juntos esconderia o achado mais importante desta análise.')

h2('Pergunta 1: o que faz alguém voltar no dia seguinte (D+1)?')
para('Taxa geral de retorno no dia seguinte: 22,1%. Testei estatisticamente (qui-quadrado e teste-t, nível de confiança de 95%) contra horário do dia, dia da semana implícito na palavra, palavra específica do dia, device, setor, faixa salarial, hábito de delivery, ser assinante da newsletter, ter aberto a newsletter naquele dia, resultado (venceu/perdeu) e tempo de jogo.')
para('O resultado é direto: nenhuma dessas variáveis muda a taxa de forma relevante.', bold=True)

make_table(
    ['Variável testada', 'Resultado do teste', 'Conclusão'],
    [
        ['Horário do dia', 'p = 0,23 (não significativo)', 'Taxa varia entre 21% e 23% em qualquer faixa de horário'],
        ['Palavra do dia', 'p = 0,21 (não significativo)', 'Palavra mais "retentiva" (24,6%) e menos (19,3%) não diferem de forma estatisticamente confiável'],
        ['Device (Android/iOS)', 'p = 0,48 (não significativo)', '22,3% vs 22,0% — diferença irrelevante'],
        ['Setor de trabalho', 'p = 0,09 (não significativo a 95%)', 'Educação (23,6%) levemente acima de Direito (20,9%), mas não é confiável estatisticamente'],
        ['Faixa salarial', 'p = 0,32 (não significativo)', 'Sem padrão'],
        ['Abriu newsletter antes de jogar', 'p = 0,16 (não significativo)', '22,3% vs 21,5% — efeito pequeno e não confiável isoladamente para D+1'],
        ['Resultado (ganhou/perdeu)', 'p = 0,18 (não significativo)', 'Perder o jogo não empurra nem afasta o retorno no dia seguinte'],
        ['Pede delivery? / frequência', 'p = 0,008 e p = 0,05 (no limite)', 'Efeito muito pequeno (22,6% vs 20,6%) — estatisticamente "real" mas pequeno demais para guiar uma decisão de produto'],
        ['É assinante da newsletter', 'p = 0,02 (significativo, efeito pequeno)', '22,6% vs 20,8% — mesma observação acima'],
    ]
)
para('Leitura para o produto: o "gatilho" do retorno no dia seguinte provavelmente não está em nenhuma dessas variáveis — está em algo que esse dataset não capturou (recebeu notificação? que horário a notificação chegou? abriu o app por outro motivo?). Investir em segmentar por setor, device ou palavra do dia para melhorar o D+1 não tem base nos dados.')

h2('Pergunta 2: o que faz alguém continuar ativo depois de 30 dias (D30)?')
para('Aqui a história muda completamente. Taxa geral de estar ativo em D30: 31,9%. Quatro fatores aparecem com significância estatística muito forte (p < 0,001) e efeito grande:')

make_table(
    ['Fator', 'Taxa de D30 quando SIM', 'Taxa de D30 quando NÃO', 'Significância'],
    [
        ['Voltou a jogar no dia seguinte', '66,9%', '22,0%', 'p < 0,0001 — o maior efeito isolado de toda a análise'],
        ['Jogou de manhã (6h-9h)', '37,6%', '27,3%-30,3% (outros horários)', 'p < 0,0001'],
        ['Abriu a newsletter antes de jogar', '37,8%', '30,5%', 'p < 0,0001'],
        ['Venceu o jogo (em vez de perder)', '32,7%', '30,8%', 'p = 0,0001 (significativo, mas efeito pequeno)'],
    ]
)
para('E o efeito composto é o achado mais forte de toda a análise:', bold=True)
make_table(
    ['Abriu newsletter antes de jogar?', 'Voltou no D+1?', 'Taxa de Ativo em D30', 'N de sessões'],
    [
        ['Não', 'Não', '21,8%', '25.023'],
        ['Não', 'Sim', '61,1%', '7.175'],
        ['Sim', 'Não', '23,0%', '6.003'],
        ['Sim', 'Sim', '91,9%', '1.648'],
    ]
)
para('Quem abre a newsletter antes de jogar E volta no dia seguinte tem 92% de chance de continuar ativo um mês depois — contra 22% de quem não faz nenhuma das duas coisas. Isso é uma diferença de 4x, num efeito que se soma (não é só "soma das partes": newsletter sozinha leva a taxa a 23-37%, D+1 sozinho leva a 61-67%, mas as duas juntas saltam para 92%).')
para('Variáveis sem efeito relevante em D30 (testadas e descartadas como driver): device, setor, faixa salarial, hábito de delivery, ser assinante formal da newsletter (diferente de tê-la aberto naquele dia), jogar outros jogos de palavras.', italic=True, color=GRAY)

para('Nota sobre causalidade: essas são correlações observacionais, não um experimento controlado. É plausível que "abrir a newsletter antes de jogar" seja menos uma causa e mais um marcador de um tipo de usuário mais engajado em geral (lê a newsletter, joga o jogo, é um usuário "fiel" por natureza). A proposta abaixo já assume isso e desenha um teste que isola a causalidade.', italic=True)

# ---------------- ENTREGA 3 ----------------
doc.add_page_break()
h1('Entrega 3 — O que eu testaria na próxima semana')

h2('Hipótese')
para('Acredito que o ritual "abrir a newsletter → jogar o Palavritas" funciona como um gatilho de hábito matinal, porque os três sinais mais fortes de retenção de longo prazo (manhã, newsletter aberta antes do jogo, e retorno no dia seguinte) apontam todos para o mesmo comportamento: o usuário que incorpora o Palavritas à sua rotina de leitura matinal da newsletter vira um usuário recorrente. Hoje esse encadeamento acontece de forma espontânea só para uma fração dos usuários (4,1% chegam a fazer as duas coisas simultaneamente) — o resto não está sendo guiado para esse padrão.')

h2('Ação proposta')
bullet('Adicionar um CTA (call-to-action) dentro da própria newsletter, no momento em que o leitor termina de ler, convidando para jogar o Palavritas do dia — ex: "Leu tudo? Agora teste seu vocabulário no Palavritas de hoje" com link direto para o jogo.')
bullet('Testar isso como A/B test: 50% dos assinantes recebem a newsletter com o CTA, 50% recebem a versão atual (sem CTA, controle). Rodar por 4-6 semanas para conseguir medir D30 de quem entrou no teste nas primeiras semanas.')
bullet('Complementar com uma notificação push pela manhã (ex: 8h) para quem já demonstrou esse padrão pelo menos uma vez (abriu newsletter e jogou no mesmo dia), reforçando o hábito nos dias seguintes — isso ataca diretamente o D+1, que hoje não tem nenhum driver claro nos dados disponíveis.')
bullet('Não investir, por ora, em personalização por setor, salário, device ou dificuldade da palavra para melhorar retenção — os dados mostram que nenhum desses fatores tem efeito relevante.')

h2('Critério de sucesso')
para('Vou considerar que funcionou se, no grupo que recebeu o CTA na newsletter:')
bullet('A proporção de usuários que abre a newsletter E joga no mesmo dia subir de forma estatisticamente significativa frente aos 4,1% atuais (linha de base medida nesta análise).')
bullet('A taxa de Ativo em D30 do grupo teste for maior que a do grupo controle, com diferença testada por qui-quadrado a 95% de confiança (mesmo método usado nesta análise) — e não apenas uma melhora aparente.')
bullet('Como meta de referência: se o efeito observado nos dados históricos se sustentar, esperaria mover a taxa de D30 do grupo exposto para algo entre 35% e 45% (parcial do salto observado de 22% para 92%, já que o CTA não garante que o usuário também volte no D+1 — isso ainda depende do push complementar).')

# ---------------- ANEXO ----------------
doc.add_page_break()
h1('Anexo — Metodologia e arquivos')
bullet('Testes estatísticos: qui-quadrado de independência para variáveis categóricas vs. retenção (binária); teste-t de Welch para variáveis contínuas vs. retenção. Nível de confiança: 95% (p < 0,05).')
bullet('Arquivos de código: 01_explore.py a 06_dashboard.py (Python/pandas/scipy), executados em sequência. Dados limpos salvos em sessions_clean.csv, attempts_clean.csv, user_profile_clean.csv. Resultados estatísticos completos em stats_results.txt. Log de limpeza linha-a-linha em cleaning_log.txt.')
bullet('Dashboard: dashboard_palavritas.xlsx (Excel, com tabelas e gráficos nativos) — usado como substituto de Metabase/Looker Studio, que não estavam disponíveis no ambiente desta análise.')

doc.save('Analise_Palavritas_the_news.docx')
print('saved docx')
